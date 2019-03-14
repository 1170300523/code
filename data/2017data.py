# coding:utf-8
# 项目: 高校评测数据分析
# 作者: 李国庆
# 时间: 2019/1/23

import pandas as pd

import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

data = pd.read_excel('test.xlsx', sheet_name='2017年原始数据')
all_school_num = len(data)
figure_list = list(data.columns) 


# 将学校分类为四种,每类为办学性质相同的集合
school = ['一流大学建设高校','一流学科建设高校','一般普通本科','独立学院']
x = []
for i in range(4):
    x.append(data[data.学校类型 == school[i]])



#函数,将学校分为4类,对600+参数中的一个加入报告
def process(info,dtable):
    if(info in figure_list):
        s ,a= [],[]
        s5 = 0 
        for i in range(4):
           s.append(round(sum(x[i][info]),2))
           a.append(round(s[i]/len(x[i]),2))
           s5 += s[i]
        s.append(round(s5,2)) 
        a.append(round(s5/all_school_num,2))
        row_cells = dtable.add_row().cells
        row_cells[0].text = '求和项: '+info
        for i in range(5):
            row_cells[i+1].text = str(s[i])
        row_cells = dtable.add_row().cells
        row_cells[0].text = '平均项: '+info
        for i in range(5):
            row_cells[i+1].text = str(a[i])
    else:
        print(info+' 不在列表中')



# 对单个学校进行搜索
def  school_search(school_name):
    oneschool(school_name)
    school_list = list(data.学校名称.values)
    school_index = school_list.index(school_name)
    school_selected = data.iloc[school_index]

    document = Document()
    document.add_heading(school_name+'综合评价',0)
    document.add_picture('school_pic/'+str(school_index)+'.png')

    table = document.add_table(1,3,style="Medium Grid 1 Accent 3");
    table.cell(0,0).text = '项目'
    table.cell(0,1).text = '数据'
    table.cell(0,2).text = '排名'

    if(school_name in school_name):
        com = data.columns
        for i in com:
            if(type(data[i][5]) is not str):
                c = data[i]
                r = len(c[c>school_selected[i]]) + 1
                row = table.add_row().cells;
                row[0].text = i
                row[1].text = str(school_selected[i])
                row[2].text = str(r) + ' / ' + str(len(data))
   
    document.save('school_report/'+school_name+'综合评价'+'.docx')


def my_rank(schoolname,job):
    school_name = list(data.学校名称.values)
    school_index = school_name.index(schoolname)
    school_selected = data.iloc[school_index]
    c = data[job]
    return len(c[c>school_selected[job]])

def ranklist(schoolname,l):
    ave = 0
    for i in l:
        ave  += 1 - my_rank(schoolname,i)/all_school_num
        ave = round(ave,3)
    return ave

# 生成一个学校综合实力的 极坐标图
def oneschool(schoolname):
    k = []
    f = open('info/key_info.txt','r')
    line = f.readline()
    while(line ):
            line = line.split(' ')
            line.remove('\n')
            k.append(line)
            line = f.readline()
    f.close()
    schooldata = []
    for i in k:
        schooldata.append(ranklist(schoolname,i))
    #标签
    labels = np.array(['办学条件','教学经费  ','教师质量','国外合作','课程教学','创新创业','课余文化','就业情况  '])  
    dataLenth = 8                   #数据个数
    datak = np.array(schooldata)    #数据
    angles = np.linspace(0, 2*np.pi, dataLenth, endpoint=False)
    datak = np.append(datak, [datak[0]])
    angles = np.append(angles, [angles[0]])
    
    fig = plt.figure()
    ax = fig.add_subplot(111, polar=True)
    ax.plot(angles, datak, 'r-', linewidth=1)
    #ax.plot(angles, datak1, 'r-', linewidth=1)
    ax.set_thetagrids(angles * 180/np.pi, labels, fontproperties="SimHei")
    #ax.set_title(schoolname + "综合评价图", va='bottom', fontproperties="SimHei",fontsize = 18)
    ax.fill(angles,datak,facecolor = 'r',alpha = 0.5)
    ax.grid(True)
    
    school_name = list(data.学校名称.values)
    school_index = school_name.index(schoolname)
    plt.savefig('school_pic/'+str(school_index),dpi = 100)
    #plt.show()

# 将所有大学分为四类进行各项参数的统计
def  big_class_survey():
        k = ['办学条件','教学经费  ','教师质量','国外合作','课程教学','创新创业','课余文化','就业情况  ']
        f = open('info/1.txt','r',encoding='utf-8' )
       
        document = Document()
        document.add_heading('2017高校数据分析',0)
        document.core_properties.author = '李国庆'
        document.add_page_break();

        line = f.readline()
        chartNum = 1;
        while(line):
            line = line.replace('\n','').split(' ')
            
            chartTitle = line.pop(0)
            document.add_heading('表'+str(chartNum)+' :'+chartTitle)
            dtable = document.add_table(1, 6, style="Table Grid",autofit = True)
            
            dtable.cell(0,0).text = '值'
            for i in range(4):
                dtable.cell(0,i+1).text = school[i]
            dtable.cell(0,5).text = '总计'
            print(line)
            for i in line:
                process(i,dtable)

            line = f.readline()
            chartNum += 1


            
        f.close()
        document.save('big_class_survey/big_class_survey.docx')


big_class_survey()
# school_search('四川大学')

#print("1.学校分类总结\n")
#print("2.单个学校搜索\n")
#print("3.两个学校比较\n")
#select = int(input("请输入序号选择功能:  "))
#
#if(select == 1):
#    big_class_survey()
#    
#
#elif(select==2):
#    string = input("请输入学校名称:    ")
#    school_search(string)
#
#else:
#    str1 = input('请输入两个学校的名称')
#    str2 = input()
#    school_compare(str1,str2)






        